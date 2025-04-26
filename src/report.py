from docx import Document

def fill_template(template_path, output_path, replacements):
    # Load the document
    doc = Document(template_path)
    
    # Replace placeholders
    for paragraph in doc.paragraphs:
        for placeholder, value in replacements.items():
            if placeholder in paragraph.text:
                paragraph.text = paragraph.text.replace(placeholder, value)
    
    # Save the modified document
    doc.save(output_path)

# Example usage
template_path = 'name.docx'
output_path = 'filled_report.docx'
replacements = {
    '{name}': 'John Doe',
    '{date}': 'January 23, 2025',
}
fill_template(template_path, output_path, replacements)

import matplotlib.pyplot as plt

def create_visualization(data, image_path):
    plt.figure()
    plt.plot(data)
    plt.title('Sample Visualization')
    plt.savefig(image_path)
    plt.close()

# Example usage
data = [1, 2, 3, 4, 5]
image_path = 'visualization.png'
create_visualization(data, image_path)

def add_image_to_doc(doc_path, image_path):
    doc = Document(doc_path)
    doc.add_picture(image_path)
    doc.save(doc_path)

# Example usage
add_image_to_doc(output_path, image_path)
